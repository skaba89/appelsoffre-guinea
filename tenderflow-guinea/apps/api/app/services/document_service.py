"""TenderFlow Guinea — Document Service.

Handles upload, text extraction, and document generation.
"""
import io
import os
from datetime import datetime
from typing import Optional
from uuid import uuid4

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.models.tender import TenderDocument


async def upload_document(
    db: AsyncSession,
    tender_id: str,
    filename: str,
    content: bytes,
    content_type: str,
) -> TenderDocument:
    """Upload a document for a tender and store it in MinIO/S3.

    Falls back to local filesystem storage if MinIO is unavailable.
    """
    storage_path = f"tenders/{tender_id}/{uuid4().hex[:8]}_{filename}"

    # Try MinIO upload
    try:
        from minio import Minio
        minio_client = Minio(
            settings.MINIO_ENDPOINT,
            access_key=settings.MINIO_ACCESS_KEY,
            secret_key=settings.MINIO_SECRET_KEY,
            secure=settings.MINIO_SECURE,
        )
        # Ensure bucket exists
        if not minio_client.bucket_exists(settings.MINIO_BUCKET):
            minio_client.make_bucket(settings.MINIO_BUCKET)

        minio_client.put_object(
            settings.MINIO_BUCKET,
            storage_path,
            io.BytesIO(content),
            length=len(content),
            content_type=content_type,
        )
    except Exception:
        # Fallback: local filesystem
        local_dir = os.path.join("/tmp", "tenderflow-docs", f"tenders/{tender_id}")
        os.makedirs(local_dir, exist_ok=True)
        local_path = os.path.join(local_dir, f"{uuid4().hex[:8]}_{filename}")
        with open(local_path, "wb") as f:
            f.write(content)
        storage_path = local_path

    # Create DB record
    doc = TenderDocument(
        tender_id=tender_id,
        filename=storage_path.split("/")[-1],
        original_filename=filename,
        file_type=content_type,
        file_size=len(content),
        storage_path=storage_path,
        is_parsed=False,
    )
    db.add(doc)
    await db.flush()

    # Try to extract text immediately
    await extract_text(db, doc.id, content)

    return doc


async def extract_text(
    db: AsyncSession,
    document_id: str,
    content: Optional[bytes] = None,
) -> Optional[str]:
    """Extract text from a document (PDF, DOCX, etc.).

    Updates the TenderDocument record with the extracted text.
    """
    result = await db.execute(
        select(TenderDocument).where(TenderDocument.id == document_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        return None

    extracted_text = ""

    if content is None:
        # Try to read from storage
        try:
            from minio import Minio
            minio_client = Minio(
                settings.MINIO_ENDPOINT,
                access_key=settings.MINIO_ACCESS_KEY,
                secret_key=settings.MINIO_SECRET_KEY,
                secure=settings.MINIO_SECURE,
            )
            response = minio_client.get_object(settings.MINIO_BUCKET, doc.storage_path)
            content = response.read()
            response.close()
            response.release_conn()
        except Exception:
            # Try local filesystem
            try:
                with open(doc.storage_path, "rb") as f:
                    content = f.read()
            except Exception:
                return None

    if content:
        file_lower = doc.original_filename.lower()

        if file_lower.endswith(".pdf"):
            extracted_text = _extract_pdf_text(content)
        elif file_lower.endswith((".docx", ".doc")):
            extracted_text = _extract_docx_text(content)
        elif file_lower.endswith(".txt"):
            extracted_text = content.decode("utf-8", errors="replace")
        elif file_lower.endswith(".html"):
            extracted_text = _extract_html_text(content)
        else:
            # Try as plain text
            try:
                extracted_text = content.decode("utf-8", errors="replace")
            except Exception:
                extracted_text = ""

    doc.content_text = extracted_text
    doc.is_parsed = True
    await db.flush()

    return extracted_text


def _extract_pdf_text(content: bytes) -> str:
    """Extract text from a PDF file."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=content, filetype="pdf")
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n\n".join(text_parts)
    except ImportError:
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            return pdfminer_extract(io.BytesIO(content))
        except ImportError:
            return "[Extraction PDF non disponible — installez PyMuPDF ou pdfminer]"


def _extract_docx_text(content: bytes) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)
    except ImportError:
        return "[Extraction DOCX non disponible — installez python-docx]"


def _extract_html_text(content: bytes) -> str:
    """Extract text from HTML content."""
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(content, "html.parser")
        # Remove script and style elements
        for element in soup(["script", "style"]):
            element.decompose()
        return soup.get_text(separator="\n", strip=True)
    except ImportError:
        return "[Extraction HTML non disponible — installez beautifulsoup4]"


def generate_docx_from_template(template_data: dict, output_path: str) -> bytes:
    """Generate a DOCX document from template data.

    This creates a basic professional document with the provided data.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Add title
        title = template_data.get("title", "Document")
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        if template_data.get("reference"):
            doc.add_paragraph(f"Référence : {template_data['reference']}")
        if template_data.get("date"):
            doc.add_paragraph(f"Date : {template_data['date']}")

        doc.add_paragraph("")

        # Add content sections
        sections = template_data.get("sections", [])
        for section in sections:
            if section.get("heading"):
                doc.add_heading(section["heading"], level=section.get("level", 1))
            if section.get("content"):
                doc.add_paragraph(section["content"])

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    except ImportError:
        return "[Generation DOCX non disponible - installez python-docx]".encode("utf-8")


def generate_pdf_from_template(template_data: dict) -> bytes:
    """Generate a PDF document from template data.

    Uses ReportLab for PDF generation.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Heading1, Heading2

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=25*mm, rightMargin=25*mm, topMargin=25*mm, bottomMargin=25*mm)
        styles = getSampleStyleSheet()
        elements = []

        # Title
        title = template_data.get("title", "Document")
        elements.append(Heading1(title))
        elements.append(Spacer(1, 10))

        # Metadata
        if template_data.get("reference"):
            elements.append(Paragraph(f"<b>Référence :</b> {template_data['reference']}", styles["Normal"]))
        if template_data.get("date"):
            elements.append(Paragraph(f"<b>Date :</b> {template_data['date']}", styles["Normal"]))
        elements.append(Spacer(1, 20))

        # Sections
        for section in template_data.get("sections", []):
            if section.get("heading"):
                elements.append(Heading2(section["heading"]))
            if section.get("content"):
                elements.append(Paragraph(section["content"], styles["Normal"]))
            elements.append(Spacer(1, 10))

        doc.build(elements)
        return buffer.getvalue()

    except ImportError:
        return "[Generation PDF non disponible - installez reportlab]".encode("utf-8")
